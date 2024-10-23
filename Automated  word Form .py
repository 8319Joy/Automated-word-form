#!/usr/bin/env python
# coding: utf-8
you can create a form in a Word file using Python. To generate a form for an admission process into a Fullstack Development course in Python, you can use the python-docx library. This library allows you to create and modify Word documents programmatically. Here's how you can do it:
# # Install python-docx:

# In[1]:


pip install python-docx


# # Use the following Python code to create a basic admission form for Fullstack Development in Python:

# In[3]:


from docx import Document
from docx.shared import Pt

def create_admission_form():
    # Create a new Document
    doc = Document()

    # Add a title
    doc.add_heading('Admission Form for Fullstack Development in Python', 0)

    # Add applicant's information fields
    doc.add_paragraph('Name of Applicant:')
    doc.add_paragraph('Date of Birth:')
    doc.add_paragraph('Email:')
    doc.add_paragraph('Phone Number:')

    # Add educational background fields
    doc.add_heading('Educational Background', level=1)
    doc.add_paragraph('Highest Degree Obtained:')
    doc.add_paragraph('University/College:')
    doc.add_paragraph('Year of Graduation:')

    # Add course-specific information
    doc.add_heading('Course Information', level=1)
    doc.add_paragraph('Preferred Start Date:')
    doc.add_paragraph('Prior Programming Experience (if any):')

    # Save the document
    doc.save('Admission_Form_Fullstack_Development.docx')

if __name__ == "__main__":
    create_admission_form()

Explanation:
The script creates a new Word document and adds headings and form fields like "Name of Applicant", "Email", "Phone Number", etc.You can easily add more fields or format the document as needed.
The resulting Word file will be saved as Admission_Form_Fullstack_Development.docx in the current directory.
This is a basic template, and you can extend it with more formatting, such as bold or italic text, spacing, or even tables for more complex forms.
# In[8]:


file_path = 'C:\\Users\\PURNANGSHU ROY\\OneDrive\\Desktop\\Data\\Admission_Form_Fullstack_Development.docx'



# In[9]:


from docx import Document

def create_admission_form():
    # Create a new Document
    doc = Document()

    # Add a title
    doc.add_heading('Admission Form for Fullstack Development in Python', 0)

    # Add applicant's information fields
    doc.add_paragraph('Name of Applicant:')
    doc.add_paragraph('Date of Birth:')
    doc.add_paragraph('Email:')
    doc.add_paragraph('Phone Number:')

    # Add educational background fields
    doc.add_heading('Educational Background', level=1)
    doc.add_paragraph('Highest Degree Obtained:')
    doc.add_paragraph('University/College:')
    doc.add_paragraph('Year of Graduation:')

    # Add course-specific information
    doc.add_heading('Course Information', level=1)
    doc.add_paragraph('Preferred Start Date:')
    doc.add_paragraph('Prior Programming Experience (if any):')

    # Save the document
    file_path = r'C:\Users\PURNANGSHU ROY\OneDrive\Desktop\Data\Admission_Form_Fullstack_Development.docx'
    doc.save(file_path)
    return file_path

# Create and save the Word file
file_path = create_admission_form()
print(f"File saved at: {file_path}")


# In[ ]:




